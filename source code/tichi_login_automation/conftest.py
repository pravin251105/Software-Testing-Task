import os
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

import pytest
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def pytest_addoption(parser):
    parser.addoption("--headless", action="store_true", default=True,
                      help="Run browser in headless mode (default: True)")


def pytest_configure(config):
    config._login_test_results = []


@pytest.hookimpl(hookwrapper=True)
def pytest_runtest_makereport(item, call):
    outcome = yield
    report = outcome.get_result()
    if report.when == "call":
        item.config._login_test_results.append(
            {
                "nodeid": report.nodeid,
                "outcome": report.outcome,
                "duration": report.duration,
                "longrepr": getattr(report, "longreprtext", ""),
            }
        )


def pytest_sessionfinish(session, exitstatus):
    results = getattr(session.config, "_login_test_results", [])
    report_dir = Path(session.config.rootpath) / "reports"
    report_dir.mkdir(parents=True, exist_ok=True)

    summary = Counter(result["outcome"] for result in results)
    total = len(results)
    passed = summary.get("passed", 0)
    failed = summary.get("failed", 0)
    skipped = summary.get("skipped", 0)
    total_duration = sum(result["duration"] for result in results)

    def escape_md(value):
        return str(value).replace("|", "\\|").replace("\n", "<br>")

    lines = [
        "# Login Automation Test Execution Report",
        "",
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}",
        f"Exit status: {exitstatus}",
        "",
        "## Summary",
        "",
        "| Metric | Value |",
        "|---|---:|",
        f"| Total test cases | {total} |",
        f"| Passed | {passed} |",
        f"| Failed | {failed} |",
        f"| Skipped | {skipped} |",
        f"| Duration | {total_duration:.2f}s |",
        "",
        "## Results",
        "",
        "| Test Case | Status | Duration |",
        "|---|---|---:|",
    ]

    for result in results:
        status = result["outcome"].upper()
        lines.append(
            f"| {escape_md(result['nodeid'])} | {status} | {result['duration']:.2f}s |"
        )

    failed_results = [result for result in results if result["outcome"] == "failed"]
    if failed_results:
        lines.extend(["", "## Failure Details", ""])
        for result in failed_results:
            lines.extend(
                [
                    f"### {escape_md(result['nodeid'])}",
                    "",
                    "```",
                    result["longrepr"] or "No traceback captured.",
                    "```",
                    "",
                ]
            )

    report_path = report_dir / "execution_report.md"
    report_path.write_text("\n".join(lines), encoding="utf-8")

    word_report_path = report_dir / "execution_report.docx"
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Tichi Login Automation Test Execution Report")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(9)

    document.add_paragraph()

    document.add_heading("Summary", level=1)
    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.columns[0].width = Inches(2.6)
    summary_table.columns[1].width = Inches(3.2)
    summary_header = summary_table.rows[0].cells
    summary_header[0].text = "Metric"
    summary_header[1].text = "Value"
    for metric, value in [
        ("Total test cases", total),
        ("Passed", passed),
        ("Failed", failed),
        ("Skipped", skipped),
        ("Duration", f"{total_duration:.2f}s"),
        ("Exit status", exitstatus),
    ]:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = str(metric)
        row_cells[1].text = str(value)

    document.add_paragraph()
    document.add_heading("Results", level=1)
    result_table = document.add_table(rows=1, cols=3)
    result_table.style = "Table Grid"
    result_table.columns[0].width = Inches(4.0)
    result_table.columns[1].width = Inches(1.0)
    result_table.columns[2].width = Inches(1.0)
    result_header = result_table.rows[0].cells
    result_header[0].text = "Test Case"
    result_header[1].text = "Status"
    result_header[2].text = "Duration"
    for result in results:
        row_cells = result_table.add_row().cells
        row_cells[0].text = result["nodeid"]
        row_cells[1].text = result["outcome"].upper()
        row_cells[2].text = f"{result['duration']:.2f}s"

    if failed_results:
        document.add_paragraph()
        document.add_heading("Failure Details", level=1)
        for result in failed_results:
            document.add_heading(result["nodeid"], level=2)
            paragraph = document.add_paragraph()
            paragraph.add_run(result["longrepr"] or "No traceback captured.")

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Generated automatically by pytest")
    footer_run.italic = True
    footer_run.font.size = Pt(8.5)

    document.save(word_report_path)


@pytest.fixture(scope="function")
def driver(request):
    options = webdriver.ChromeOptions()
    if request.config.getoption("--headless"):
        options.add_argument("--headless=new")
    options.add_argument("--window-size=1366,900")
    options.add_argument("--disable-gpu")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")

    # webdriver-manager can occasionally return a non-executable metadata file path
    # on Windows; normalize to the actual chromedriver binary.
    driver_path = Path(ChromeDriverManager().install())
    if (
        not driver_path.name.lower().startswith("chromedriver")
        or driver_path.suffix.lower() not in {"", ".exe"}
    ):
        candidate = next(
            (p for p in driver_path.parent.glob("chromedriver*.exe") if p.is_file()),
            None,
        )
        if candidate is not None:
            driver_path = candidate

    service = Service(str(driver_path))
    drv = webdriver.Chrome(service=service, options=options)
    drv.implicitly_wait(5)

    yield drv

    drv.quit()


@pytest.fixture(scope="session")
def valid_credentials():
    """
    Test account credentials, supplied via environment variables so real
    credentials are never hard-coded / committed to source control.

    Set before running:
        export TICHI_TEST_EMAIL="youraccount@example.com"
        export TICHI_TEST_PASSWORD="YourPassword123"
    """
    return {
        "email": os.getenv("TICHI_TEST_EMAIL", "REPLACE_WITH_VALID_TEST_EMAIL"),
        "password": os.getenv("TICHI_TEST_PASSWORD", "REPLACE_WITH_VALID_TEST_PASSWORD"),
    }
